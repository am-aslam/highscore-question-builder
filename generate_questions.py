import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt

OUT_DIR = Path.cwd() / "highscore_output"
OUT_DIR.mkdir(exist_ok=True)

SAMPLE_TEXT = """@title Algebra & Geometry Mini Assessment
@description A mini assessment containing two AI-generated math multiple-choice questions similar in style to the base questions. Each question preserves LaTeX where appropriate and follows the requested output format.

@question If $10x - 5 = 5x + 15$, what is the value of x?
@instruction Solve for x and choose the correct answer.
@difficulty easy
@Order 1

@option \\(\\frac{2}{3}\\)
@option \\(\\frac{3}{2}\\)
@@option Correct Answer
@option 4
@option 5

@explanation Subtract 5x from both sides: 5x - 5 = 15 -> 5x = 20 -> x = 4.
@subject Quantitative Math
@unit Algebra
@topic Interpreting Variables

@plusmarks 1


@question A right triangle has legs of length 6 and 8. What is the measure of the hypotenuse?
@instruction Select the exact length.
@difficulty easy
@Order 2

@option 10
@@option Correct Answer
@option 12
@option 14

@explanation Use Pythagorean theorem: hypotenuse = sqrt(6^2 + 8^2) = sqrt(100) = 10.
@subject Quantitative Math
@unit Geometry and Measurement
@topic Right Triangles & Trigonometry

@plusmarks 1
"""

def parse_questions(text: str):
    """Parses raw @question format into a list of structured dictionaries."""
    questions = []
    blocks = text.strip().split("@question")[1:]

    for block in blocks:
        q_data = {}
        lines = block.strip().splitlines()
        q_data["question"] = lines[0].strip()
        for line in lines[1:]:
            if line.startswith("@instruction"):
                q_data["instruction"] = line.replace("@instruction", "").strip()
            elif line.startswith("@difficulty"):
                q_data["difficulty"] = line.replace("@difficulty", "").strip()
            elif line.startswith("@Order"):
                q_data["order"] = line.replace("@Order", "").strip()
            elif line.startswith("@@option"):
                q_data["correct"] = line.replace("@@option", "").strip()
            elif line.startswith("@option"):
                q_data.setdefault("options", []).append(line.replace("@option", "").strip())
            elif line.startswith("@explanation"):
                q_data["explanation"] = line.replace("@explanation", "").strip()
            elif line.startswith("@subject"):
                q_data["subject"] = line.replace("@subject", "").strip()
            elif line.startswith("@unit"):
                q_data["unit"] = line.replace("@unit", "").strip()
            elif line.startswith("@topic"):
                q_data["topic"] = line.replace("@topic", "").strip()
            elif line.startswith("@plusmarks"):
                q_data["marks"] = line.replace("@plusmarks", "").strip()
        questions.append(q_data)
    return questions

def write_professional_docx(text: str, filepath: Path):
    doc = Document()
    doc.add_heading("HighScore.ai Assignment – AI-Generated Mathematics Questions", level=1)
    doc.add_paragraph("Total Marks: 2")
    doc.add_paragraph("Instructions: Answer the following questions. Each question carries 1 mark.\n")

    questions = parse_questions(text)
    for q in questions:
        doc.add_paragraph(f"Q{q['order']}. ({q['difficulty'].capitalize()}) {q['unit']} – {q['topic']} ({q['marks']} mark{'s' if q['marks'] != '1' else ''})", style='List Number')
        doc.add_paragraph(q['question'])
        if "instruction" in q:
            doc.add_paragraph(f"Instruction: {q['instruction']}")

        for opt in q.get("options", []):
            prefix = "✔ " if opt == q.get("correct", "") else "• "
            doc.add_paragraph(f"{prefix}{opt}", style='List Bullet')

        doc.add_paragraph("Explanation:", style='Intense Quote')
        doc.add_paragraph(q.get("explanation", "No explanation provided."))
        doc.add_paragraph() 

    filepath.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(filepath))
    print(f"Saved professional .docx to: {filepath}")

def write_text_file(text: str, filepath: Path):
    filepath.parent.mkdir(parents=True, exist_ok=True)
    filepath.write_text(text, encoding="utf-8")
    print(f"Saved raw text to: {filepath}")

def generate_with_gemini(prompt: str, outfile: str = "gemini_generated.txt"):
    try:
        import google.generativeai as genai
    except ImportError:
        raise RuntimeError("google-generativeai package not installed. Run: pip install google-generativeai")

    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        raise RuntimeError("GEMINI_API_KEY environment variable not set. Export it and retry.")

    genai.configure(api_key=api_key)
    model = genai.GenerativeModel("gemini-1.5-flash")
    response = model.generate_content(prompt)
    text = response.text.strip()

    outpath = OUT_DIR / outfile
    write_text_file(text, outpath)
    write_professional_docx(text, OUT_DIR / (outfile.replace(".txt", ".docx")))
    return text

PROMPT_TEMPLATE = """
You are an assistant that generates multiple-choice math questions.
Requirements:
- Use LaTeX for equations where needed (keep $...$).
- Follow the exact Question Output Format shown below.
- Choose subject, unit, topic from the provided curriculum list.
- Provide one correct answer marked with '@@option Correct Answer' and include an explanation.

Template:
@question [QUESTION_TEXT]
@instruction [INSTRUCTION]
@difficulty [easy/moderate/hard]
@Order [NUMBER]

@option [A]
@option [B]
@@option Correct Answer
@option [D]
@explanation [EXPLANATION]
@subject [subject]
@unit [unit]
@topic [topic]
@plusmarks 1
"""

def main():
    txt_path = OUT_DIR / "generated_questions_sample.txt"
    docx_path = OUT_DIR / "generated_questions_sample.docx"
    write_text_file(SAMPLE_TEXT, txt_path)
    write_professional_docx(SAMPLE_TEXT, docx_path)

    print("\nSample files created at:", OUT_DIR)
    print("To generate new questions with Gemini, run:")
    print("   python generate_questions.py gemini")

if __name__ == "__main__":
    if len(sys.argv) > 1:
        mode = sys.argv[1].lower()
        prompt = PROMPT_TEMPLATE + "\nNow create TWO new MCQ questions similar to the base question: 'Six equally-sized pieces are cut from a piece of rope that has a length of 88 inches. Each of the six pieces has a length of 14 inches. What is the length, in inches, of the leftover piece?' Provide outputs in exact Question Output Format.\n"
        if mode == "gemini":
            print("Generating via Gemini API...")
            try:
                out = generate_with_gemini(prompt, outfile="gemini_generated.txt")
                print("Gemini output saved and formatted.")
            except Exception as e:
                print("Error:", e)
                sys.exit(1)
        else:
            print("Unknown mode. Use 'gemini'.")
            sys.exit(1)
    else:
        main()
