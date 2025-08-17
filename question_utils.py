# question_utils.py

import os
from pathlib import Path
from docx import Document

OUT_DIR = Path.cwd() / "highscore_output"
OUT_DIR.mkdir(exist_ok=True)

def parse_questions(text: str):
    questions = []
    blocks = text.strip().split("@question")[1:]
    for block in blocks:
        q = {}
        lines = block.strip().splitlines()
        q["question"] = lines[0].strip()
        for line in lines[1:]:
            if line.startswith("@instruction"):
                q["instruction"] = line.replace("@instruction", "").strip()
            elif line.startswith("@difficulty"):
                q["difficulty"] = line.replace("@difficulty", "").strip()
            elif line.startswith("@Order"):
                q["order"] = line.replace("@Order", "").strip()
            elif line.startswith("@@option"):
                q["correct"] = line.replace("@@option", "").strip()
            elif line.startswith("@option"):
                q.setdefault("options", []).append(line.replace("@option", "").strip())
            elif line.startswith("@explanation"):
                q["explanation"] = line.replace("@explanation", "").strip()
            elif line.startswith("@subject"):
                q["subject"] = line.replace("@subject", "").strip()
            elif line.startswith("@unit"):
                q["unit"] = line.replace("@unit", "").strip()
            elif line.startswith("@topic"):
                q["topic"] = line.replace("@topic", "").strip()
            elif line.startswith("@plusmarks"):
                q["marks"] = line.replace("@plusmarks", "").strip()
        questions.append(q)
    return questions

def write_docx(text: str, filepath: Path):
    doc = Document()
    doc.add_heading("HighScore.ai Assignment – AI-Generated Math Questions", level=1)
    doc.add_paragraph("Each question carries 1 mark.\n")

    questions = parse_questions(text)
    for q in questions:
        doc.add_paragraph(f"Q{q['order']}. {q['question']} ({q.get('marks', '1')} mark)", style='List Number')
        for opt in q.get("options", []):
            prefix = "✔ " if opt == q.get("correct", "") else "• "
            doc.add_paragraph(f"{prefix}{opt}", style='List Bullet')
        doc.add_paragraph("Explanation: " + q.get("explanation", ""))
        doc.add_paragraph()
    doc.save(filepath)
    return filepath

def write_text_file(text: str, filepath: Path):
    filepath.parent.mkdir(parents=True, exist_ok=True)
    filepath.write_text(text, encoding="utf-8")
    return filepath

def generate_with_gemini(prompt: str, outfile: str = "gemini_generated.txt"):
    import google.generativeai as genai

    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        raise RuntimeError("GEMINI_API_KEY not set")

    genai.configure(api_key=api_key)
    model = genai.GenerativeModel("gemini-1.5-flash")
    response = model.generate_content(prompt)
    text = response.text.strip()

    outpath = OUT_DIR / outfile
    write_text_file(text, outpath)
    write_docx(text, OUT_DIR / (outfile.replace(".txt", ".docx")))
    return text
