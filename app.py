import streamlit as st
from pathlib import Path
from docx import Document
import os

OUT_DIR = Path.cwd() / "highscore_output"
OUT_DIR.mkdir(exist_ok=True)

SAMPLE_TEXT = """@question If $10x - 5 = 5x + 15$, what is the value of x?
@instruction Solve for x and choose the correct answer.
@difficulty easy
@Order 1

@option \\(\\frac{2}{3}\\)
@option \\(\\frac{3}{2}\\)
@@option Correct Answer
@option 4
@option 5

@explanation Subtract 5x from both sides: 5x - 5 = 15 → 5x = 20 → x = 4.
@subject Quantitative Math
@unit Algebra
@topic Interpreting Variables
@plusmarks 1
"""

# 1. Function to parse @question text
def parse_questions(text: str):
    questions = []
    blocks = text.strip().split("@question")[1:]

    for block in blocks:
        q = {}
        lines = block.strip().splitlines()
        q["question"] = lines[0].strip()
        for line in lines[1:]:
            if line.startswith("@instruction"):
                q["instruction"] = line.replace("@instruction", "").strip()
            elif line.startswith("@difficulty"):
                q["difficulty"] = line.replace("@difficulty", "").strip()
            elif line.startswith("@Order"):
                q["order"] = line.replace("@Order", "").strip()
            elif line.startswith("@@option"):
                q["correct"] = line.replace("@@option", "").strip()
            elif line.startswith("@option"):
                q.setdefault("options", []).append(line.replace("@option", "").strip())
            elif line.startswith("@explanation"):
                q["explanation"] = line.replace("@explanation", "").strip()
            elif line.startswith("@subject"):
                q["subject"] = line.replace("@subject", "").strip()
            elif line.startswith("@unit"):
                q["unit"] = line.replace("@unit", "").strip()
            elif line.startswith("@topic"):
                q["topic"] = line.replace("@topic", "").strip()
            elif line.startswith("@plusmarks"):
                q["marks"] = line.replace("@plusmarks", "").strip()
        questions.append(q)
    return questions

# 2. Generate .docx from question format
def write_docx(text: str, filepath: Path):
    doc = Document()
    doc.add_heading("HighScore.ai Assignment – AI-Generated Math Questions", level=1)
    doc.add_paragraph("Each question carries 1 mark.\n")

    questions = parse_questions(text)
    for q in questions:
        doc.add_paragraph(f"Q{q['order']}. {q['question']} ({q.get('marks', '1')} mark)", style='List Number')
        for opt in q.get("options", []):
            prefix = "✔ " if opt == q.get("correct", "") else "• "
            doc.add_paragraph(f"{prefix}{opt}", style='List Bullet')
        doc.add_paragraph("Explanation: " + q.get("explanation", ""))
        doc.add_paragraph()
    doc.save(filepath)
    return filepath

# 3. Streamlit App UI
st.title("📘 HighScore.ai – MCQ Generator")
st.markdown("Generate multiple-choice math questions and download them as `.docx`.")

# Text Area
text_input = st.text_area("📝 Input Question Format", SAMPLE_TEXT, height=300)

# Buttons
col1, col2 = st.columns(2)

with col1:
    if st.button("📄 Generate DOCX"):
        docx_path = OUT_DIR / "questions_output.docx"
        write_docx(text_input, docx_path)
        with open(docx_path, "rb") as f:
            st.download_button("⬇️ Download DOCX", f, file_name="HighScore_Questions.docx")

with col2:
    if st.button("💾 Save Text File"):
        text_path = OUT_DIR / "questions_output.txt"
        text_path.write_text(text_input, encoding="utf-8")
        with open(text_path, "rb") as f:
            st.download_button("⬇️ Download TXT", f, file_name="HighScore_Questions.txt")
